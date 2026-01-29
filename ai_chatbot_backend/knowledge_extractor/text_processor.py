# knowledge_extractor/text_processor.py

import fitz  # PyMuPDF
import re
import spacy
from pathlib import Path
from typing import List, Optional

# 全局变量用于存储SpaCy模型，避免重复加载
_nlp = None


def load_spacy_model(model_name="zh_core_web_sm"):
    """加载SpaCy模型，如果已加载则直接返回"""
    global _nlp
    if _nlp is None:
        try:
            _nlp = spacy.load(model_name)
            print(f"SpaCy model '{model_name}' loaded successfully.")
        except OSError:
            print(f"SpaCy model '{model_name}' not found. Downloading...")
            spacy.cli.download(model_name)
            _nlp = spacy.load(model_name)
    return _nlp


def extract_text_from_pdf(pdf_path: Path) -> Optional[str]:
    """从PDF文件中提取文本"""
    text = ""
    try:
        doc = fitz.open(pdf_path)
        for page_num in range(doc.page_count):
            page = doc.load_page(page_num)
            text += page.get_text()
        doc.close()
        return text
    except Exception as e:
        print(f"Error extracting text from PDF {pdf_path}: {e}")
        return None


def extract_text_from_docx(docx_path: Path) -> Optional[str]:
    """
    从DOCX文件中提取文本，包含段落和表格内容。
    ✅ 重点修改：增加了表格提取逻辑，将表格行转换为句子。
    """
    try:
        from docx import Document
        document = Document(docx_path)
        full_text = []

        # 1. 提取普通段落
        for para in document.paragraphs:
            if para.text.strip():
                full_text.append(para.text.strip())

        # 2. 提取表格内容
        for table in document.tables:
            for row in table.rows:
                # 获取该行所有单元格的文本，去除空格
                row_cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_cells:
                    # 将这一行的单元格用空格连接，形成一个类似句子的字符串
                    # 例如： "信息技术系 数字媒体技术 20000"
                    full_text.append(" ".join(row_cells))

        return '\n'.join(full_text)
    except Exception as e:
        print(f"Error extracting text from DOCX {docx_path}: {e}")
        return None


def read_text_from_txt(txt_path: Path) -> Optional[str]:
    """从TXT文件中读取文本"""
    try:
        with open(txt_path, 'r', encoding='utf-8') as f:
            return f.read()
    except Exception as e:
        print(f"Error reading text from TXT {txt_path}: {e}")
        return None


def get_text_from_file(filepath: Path) -> Optional[str]:
    """根据文件类型调用相应的文本提取函数"""
    if not filepath.exists():
        print(f"Warning: File '{filepath}' does not exist.")
        return None

    if filepath.suffix == '.pdf':
        return extract_text_from_pdf(filepath)
    elif filepath.suffix == '.docx':
        return extract_text_from_docx(filepath)
    elif filepath.suffix in ['.txt', '.md']:
        return read_text_from_txt(filepath)
    else:
        print(f"Unsupported file type for text extraction: {filepath.suffix}")
        return None


def clean_text(text: str) -> str:
    """清洗文本，去除噪音、多余空格和特殊字符"""
    if not text:
        return ""
    # 替换各种空格为标准空格
    text = re.sub(r'[\s\u3000]+', ' ', text)
    # 去除常见的页眉页脚（需要根据实际文档内容调整正则表达式）
    text = re.sub(r'\d{4}年\d{1,2}月\d{1,2}日', '', text)  # 示例：日期
    text = re.sub(r'第\s*\d+\s*页', '', text)  # 示例：页码
    # 去除一些不常见的控制字符或符号，但保留关键符号如小数点
    text = re.sub(r'[^\u4e00-\u9fa5a-zA-Z0-9\s.,;!?:()（）《》“”-]', '', text)
    return text.strip()


def segment_sentences(text: str) -> List[str]:
    """使用SpaCy进行分句"""
    nlp_model = load_spacy_model()
    # 增加最大长度限制，防止超长文本卡死
    nlp_model.max_length = 2000000
    doc = nlp_model(text)
    # 过滤掉空句子
    return [sent.text.strip() for sent in doc.sents if sent.text.strip()]